from docx import Document
import json


surgeryDict = {
    "chase jeanne": {
        "datesList": [
            "4292019"
        ],
        "pagesList": [
            1
        ]
    },
    "sas/no pes": {
        "datesList": [
            "8-19-1971"
        ],
        "pagesList": [
            10
        ]
    },
    "removal of maxillary plates with narrowing of the alar base": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            22
        ]
    },
    "radiofrequency reduction of the palate": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            22
        ]
    },
    "cei": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            25
        ]
    },
    "cfi": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            25
        ]
    },
    "cei cfi page 483": {
        "datesList": [
            "11/01/11"
        ],
        "pagesList": [
            37
        ]
    },
    "implant/explant": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            43
        ]
    },
    "hip surgery": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            52,
            87
        ]
    },
    "tonsillectomy": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            52
        ]
    },
    "mri scan/arthrogram of the right hip capsule": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            66
        ]
    },
    "right hip surgery": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            70
        ]
    },
    "both legs": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            85
        ]
    },
    "righf +arsl": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            85
        ]
    },
    "left + arsl": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            85
        ]
    },
    "hysterectomy": {
        "datesList": [
            "None"
        ],
        "pagesList": [
            87
        ]
    },
    "encounter for screening for malignant neoplasm of cervix": {
        "datesList": [
            "05/25/2021"
        ],
        "pagesList": [
            117
        ]
    }
}
medicationDict = {
        "chase jeanne": {
        "datesList": [
            [
                "4/29/2019",
                "None"
            ]
        ],
        "pagesList": [
            1
        ]
    },
    "sleep medicine center": {
        "datesList": [
            [
                "8/13/2019",
                "None"
            ]
        ],
        "pagesList": [
            4
        ]
    },
    "estradiol": {
        "datesList": [
            [
                "None",
                "None"
            ],
            [
                "09/26/2019",
                "02/23/2020"
            ]
        ],
        "pagesList": [
            6,
            109,
            110
        ]
    },
    "fluticasone propionate": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            6
        ]
    },
    "lory na (28)": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            6
        ]
    },
    "quetiapine": {
        "datesList": [
            [
                "None",
                "None"
            ],
            [
                "02/06/2023 07:45",
                "08/05/2023 07:45"
            ],
            [
                "07/02/2019",
                "None"
            ],
            [
                "11/21/2018 03:09",
                "04/09/2019 01:32 PM"
            ]
        ],
        "pagesList": [
            6,
            95,
            109,
            111
        ]
    },
    "zolpidem": {
        "datesList": [
            [
                "None",
                "None"
            ],
            [
                "06/06/2023 07:45",
                "None"
            ],
            [
                "09/26/2019",
                "02/23/2020"
            ],
            [
                "08:52 08/03/2018 03:16 PM",
                "12:00 08/03/2018 03:16 PM"
            ]
        ],
        "pagesList": [
            6,
            95,
            109,
            111
        ]
    },
    "temazepam": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            22,
            52
        ]
    },
    "trazodone": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            22,
            52
        ]
    },
    "cymbalta": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            22,
            52
        ]
    },
    "cei": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            25
        ]
    },
    "cfi": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            25
        ]
    },
    "cei cfi": {
        "datesList": [
            [
                "12/12/11",
                "None"
            ]
        ],
        "pagesList": [
            35
        ]
    },
    "chlorhexidine scrubs": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "cefazolin (kefzol)": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "cefoxitin": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "ertapenem": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "gentamicin": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "clindamycin": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "vancomycin": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "sequential compression device (scd)": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            37
        ]
    },
    "mobic": {
        "datesList": [
            [
                "11/30/2004",
                "None"
            ],
            [
                "11/30/2004",
                "None"
            ]
        ],
        "pagesList": [
            60,
            63
        ]
    },
    "15-1664": {
        "datesList": [
            [
                "Jan-00",
                "2"
            ],
            [
                "Jan-00",
                "4"
            ]
        ],
        "pagesList": [
            82,
            84
        ]
    },
    "guaifenesin/pseudoephedrne hci": {
        "datesList": [
            [
                "05/15/23",
                "None"
            ]
        ],
        "pagesList": [
            87
        ]
    },
    "oxymetazoline hci": {
        "datesList": [
            [
                "05/15/23",
                "None"
            ]
        ],
        "pagesList": [
            87
        ]
    },
    "citalopram 40 mg oral tablet": {
        "datesList": [
            [
                "05/17/2023",
                "05/15/2023"
            ],
            [
                "04/12/2022 02:32",
                "09/09/2022 02:32 PM"
            ],
            [
                "01/24/2022",
                "None"
            ],
            [
                "09/09/2021",
                "03/08/2022"
            ]
        ],
        "pagesList": [
            94,
            98,
            99,
            100
        ]
    },
    "zolpidem 10 mg oral tablet": {
        "datesList": [
            [
                "05/15/2023",
                "None"
            ],
            [
                "10/19/2022",
                "02/16/2023"
            ],
            [
                "04/12/2022 02:32",
                "05/12/2022 02:32 PM"
            ],
            [
                "01/24/2022",
                "None"
            ],
            [
                "07/27/2021",
                "None"
            ],
            [
                "07/13/2021",
                "01/09/2022"
            ],
            [
                "06/01/2021",
                "10/29/2021"
            ],
            [
                "08/02/21",
                "None"
            ],
            [
                "01/05/2021 03:48",
                "06/04/2021 03:48"
            ],
            [
                "09/28/2020",
                "04/02/2021"
            ],
            [
                "08/06/2020",
                "01/03/2021"
            ]
        ],
        "pagesList": [
            94,
            97,
            98,
            99,
            101,
            101,
            102,
            103,
            104,
            105,
            106
        ]
    },
    "methylphenidate 10 mg oral tablet": {
        "datesList": [
            [
                "05/10/2023",
                "05/10/2024"
            ],
            [
                "11/02/2022",
                "completed"
            ],
            [
                "07/05/2022",
                "08/04/2022"
            ],
            [
                "04/20/2022",
                "None"
            ],
            [
                "01/24/2022",
                "06/23/2022"
            ]
        ],
        "pagesList": [
            94,
            96,
            97,
            98,
            99
        ]
    },
    "quetiapine 50 mg oral tablet": {
        "datesList": [
            [
                "05/08/2023",
                "06/07/2023"
            ],
            [
                "12/20/2022",
                "01/05/2023"
            ],
            [
                "04/12/2022 02:32",
                "10/09/2022 02:32 PM"
            ],
            [
                "01/24/2022",
                "03/24/2022"
            ],
            [
                "07/13/2021",
                "12/10/2021"
            ],
            [
                "05/05/2021",
                "11/01/2021"
            ],
            [
                "03/05/2021",
                "08/02/2021"
            ],
            [
                "01/05/2021 03:48",
                "06/04/2021 03:48"
            ],
            [
                "11/03/2020 12:06",
                "04/02/2021 12:06"
            ],
            [
                "09/21/2020",
                "02/18/2021"
            ],
            [
                "08/06/2020",
                "01/03/2021"
            ],
            [
                "03/10/2020 02:30",
                "08/07/2020 PM 02:30"
            ]
        ],
        "pagesList": [
            94,
            96,
            98,
            99,
            101,
            102,
            103,
            104,
            104,
            105,
            106,
            108
        ]
    },
    "hydroxyzine": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            94
        ]
    },
    "citalopram": {
        "datesList": [
            [
                "02/06/2023 07:45",
                "11/04/2023 12:53"
            ],
            [
                "08/09/2019",
                "11/02/2022"
            ]
        ],
        "pagesList": [
            95,
            109
        ]
    },
    "azithromycin": {
        "datesList": [
            [
                "03/20/2023",
                "03/23/2023"
            ]
        ],
        "pagesList": [
            95
        ]
    },
    "methylphenidate": {
        "datesList": [
            [
                "02/27/2023",
                "03/23/2023"
            ],
            [
                "Unknown",
                "None"
            ]
        ],
        "pagesList": [
            95,
            109
        ]
    },
    "voltaren topical 1% topical gel": {
        "datesList": [
            [
                "06/23/2022",
                "None"
            ],
            [
                "07/27/2021",
                "07/27/2021"
            ],
            [
                "01/11/2021",
                "None"
            ]
        ],
        "pagesList": [
            97,
            101,
            103
        ]
    },
    "fluticasone nasal 0.05 mg/inh spray": {
        "datesList": [
            [
                "07/11/2022",
                "None"
            ],
            [
                "07/27/2021",
                "07/27/2021"
            ],
            [
                "06/11/2020",
                "None"
            ],
            [
                "11/20/2019",
                "11/21/2019"
            ]
        ],
        "pagesList": [
            97,
            101,
            106,
            108
        ]
    },
    "atorvastatin 20 mg oral tablet": {
        "datesList": [
            [
                "07/11/2022",
                "None"
            ]
        ],
        "pagesList": [
            97
        ]
    },
    "gabapentin": {
        "datesList": [
            [
                "05/23/2022",
                "05/23/2022"
            ]
        ],
        "pagesList": [
            98
        ]
    },
    "premarin vaginal 0.625 mg/gm cream": {
        "datesList": [
            [
                "04/27/2022",
                "04/27/2022"
            ],
            [
                "01/24/2022",
                "None"
            ]
        ],
        "pagesList": [
            98,
            99
        ]
    },
    "clindamycin 300 mg oral capsule": {
        "datesList": [
            [
                "02/22/2022",
                "03/24/2022"
            ]
        ],
        "pagesList": [
            99
        ]
    },
    "gabapentin 300 mg oral capsule": {
        "datesList": [
            [
                "01/24/2022",
                "07/23/2022"
            ]
        ],
        "pagesList": [
            99
        ]
    },
    "tramadol 50 mg oral tablet": {
        "datesList": [
            [
                "11/17/2021",
                "01/17/2022"
            ]
        ],
        "pagesList": [
            100
        ]
    },
    "citalopram 20 mg oral tablet": {
        "datesList": [
            [
                "09/09/2021",
                "02/06/2022"
            ],
            [
                "07/23/2021",
                "12/20/2021"
            ],
            [
                "03/05/2021",
                "08/02/2021"
            ],
            [
                "01/05/2021 03:48",
                "06/04/2021 03:48"
            ],
            [
                "11/03/2020 12:06",
                "03/03/2021 12:06"
            ],
            [
                "09/21/2020",
                "02/18/2021"
            ],
            [
                "07/09/2020",
                "None"
            ]
        ],
        "pagesList": [
            100,
            101,
            103,
            104,
            104,
            105,
            106
        ]
    },
    "tretinoin": {
        "datesList": [
            [
                "07/27/2021",
                "07/27/2021"
            ]
        ],
        "pagesList": [
            101
        ]
    },
    "tizanidine 4 mg oral tablet": {
        "datesList": [
            [
                "07/27/2021",
                "07/27/2021"
            ],
            [
                "08/20/2020",
                "None"
            ]
        ],
        "pagesList": [
            101,
            105
        ]
    },
    "atorvastatin 10 mg oral tablet": {
        "datesList": [
            [
                "05/26/2021",
                "09/22/2021"
            ]
        ],
        "pagesList": [
            102,
            106,
            110
        ]
    },
    "diflucan 150 mg oral tablet": {
        "datesList": [
            [
                "03/23/2021",
                "None"
            ]
        ],
        "pagesList": [
            103
        ]
    },
    "medroxyprogesterone 5 mg oral tablet": {
        "datesList": [
            [
                "01/12/2021",
                "None"
            ]
        ],
        "pagesList": [
            103
        ]
    },
    "estradiol 1 mg oral tablet": {
        "datesList": [
            [
                "08/25/2020",
                "None"
            ]
        ],
        "pagesList": [
            105
        ]
    },
    "loratadine 10 mg oral tablet": {
        "datesList": [
            [
                "05/06/2020",
                "10/03/2020"
            ]
        ],
        "pagesList": [
            106
        ]
    },
    "tizanidine": {
        "datesList": [
            [
                "09/25/2019",
                "09/25/2019"
            ],
            [
                "01:00 04/09/2019 01:31 PM",
                "03:03 03/31/2021 05:59 PM"
            ]
        ],
        "pagesList": [
            109,
            111
        ]
    },
    "fluticasone": {
        "datesList": [
            [
                "07/25/2019",
                "None"
            ]
        ],
        "pagesList": [
            109
        ]
    },
    "loratadine": {
        "datesList": [
            [
                "07/03/2019",
                "None"
            ]
        ],
        "pagesList": [
            109
        ]
    },
    "atorvastatin": {
        "datesList": [
            [
                "07/02/2019",
                "03/23/2023"
            ]
        ],
        "pagesList": [
            109
        ]
    },
    "loratadine (claritin) 10 mg tab": {
        "datesList": [
            [
                "05/31/2019",
                "04/09/2019"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "guaifenesin (mucinex) 600 mg tab": {
        "datesList": [
            [
                "completed",
                "completed"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "quetiapine (seroquel) 50 mg tablet": {
        "datesList": [
            [
                "05/08/2019",
                "active 07/02/2019"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "quetiapine (seroquel) 50 mg tab": {
        "datesList": [
            [
                "03/31/2021",
                "completed"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "ambien 10 mg oral tablet": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "celexa 20 mg oral tablet": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "estrace 1 mg oral tablet": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "seroquel 50 mg oral tablet": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            110
        ]
    },
    "tizanidine (zanaflex) 4 mg tab": {
        "datesList": [
            [
                "None",
                "None"
            ],
            [
                "05/2/2012",
                "06/30/2017"
            ]
        ],
        "pagesList": [
            110,
            114
        ]
    },
    "duloxetine hcl": {
        "datesList": [
            [
                "02:59 11/20/2018 02:52 PM",
                "12:09 07/03/2019 08:29 AM"
            ]
        ],
        "pagesList": [
            111
        ]
    },
    "citalopram (celexa) 20 mg tab": {
        "datesList": [
            [
                "10/19/2017 02:43",
                "12/28/2017 05:22 PM"
            ],
            [
                "08/05/2017 09:39",
                "10/19/2017 02:42 PM"
            ]
        ],
        "pagesList": [
            113,
            113
        ]
    },
    "bupropion sr (wellbutrin sr) 150 mg tab": {
        "datesList": [
            [
                "TEER; 02:42 04/10/2018",
                "TROT 02:42 09/18/2018"
            ]
        ],
        "pagesList": [
            113
        ]
    },
    "cefuroxime (ceftin) 250 mg tab": {
        "datesList": [
            [
                "05/30/2017 12:00",
                "05/22/2018 01:26 PM"
            ]
        ],
        "pagesList": [
            113
        ]
    },
    "trazodone hcl (trazodone po)": {
        "datesList": [
            [
                "02/10/2015",
                "05/26/2017"
            ]
        ],
        "pagesList": [
            114
        ]
    },
    "duloxetine hcl (cymbalta po)": {
        "datesList": [
            [
                "05/26/2017",
                "09/18/2018"
            ]
        ],
        "pagesList": [
            114
        ]
    },
    "temazepam po": {
        "datesList": [
            [
                "05/2/2012",
                "05/26/2017"
            ]
        ],
        "pagesList": [
            114
        ]
    },
    "zolpidem (ambien) 10 mg tab": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            114
        ]
    },
    "gabapentin (neurontin) 100 mg cap": {
        "datesList": [
            [
                "None",
                "None"
            ]
        ],
        "pagesList": [
            114
        ]
    },
    "hydrocodone-acetaminophen (vicodin)": {
        "datesList": [
            [
                "11/01/2011 11:30 PM",
                "11/02/2011 03:38 PM"
            ]
        ],
        "pagesList": [
            115
        ]
    },
    "celexa 40 mg tablet (citalopram 40 mg\noral tablet [celexa])": {
        "datesList": [
            [
                "11/01/2011",
                "11/03/2011"
            ]
        ],
        "pagesList": [
            116
        ]
    },
    "flexeril lomg oral": {
        "datesList": [
            [
                "11/01/2011",
                "11/01/2011"
            ]
        ],
        "pagesList": [
            116
        ]
    },
    "rash and other nonspecific skin eruption": {
        "datesList": [
            [
                "05/25/2021",
                "None"
            ]
        ],
        "pagesList": [
            117
        ]
    },
    "major depressive disorder, single episode, unspecified": {
        "datesList": [
            [
                "05/25/2021",
                "None"
            ]
        ],
        "pagesList": [
            117
        ]
    }
}
allergyDict = {
    "penicillin, causing rash": {
        "pagesList": [
            22
        ]
    },
    "cei": {
        "pagesList": [
            25
        ]
    },
    "cfi": {
        "pagesList": [
            25
        ]
    },
    "b/l": {
        "pagesList": [
            35
        ]
    },
    "v/e": {
        "pagesList": [
            35
        ]
    },
    "b": {
        "pagesList": [
            35
        ]
    },
    "ohee": {
        "pagesList": [
            35
        ]
    },
    "z": {
        "pagesList": [
            35
        ]
    },
    "p/e": {
        "pagesList": [
            35
        ]
    },
    "iohez": {
        "pagesList": [
            35
        ]
    },
    "penicillin": {
        "pagesList": [
            52,
            72,
            76
        ]
    },
    "wheat": {
        "pagesList": [
            72,
            76
        ]
    },
    "soy": {
        "pagesList": [
            72,
            76
        ]
    },
    "penicillins": {
        "pagesList": [
            87
        ]
    },
    "medication": {
        "pagesList": [
            94,
            98,
            101,
            102,
            105,
            107,
            109,
            110,
            111,
            112,
            113,
            114
        ]
    },
    "non-celiac gluten sensitivity": {
        "pagesList": [
            117
        ]
    },
    "allergic rhinitis": {
        "pagesList": [
            117
        ]
    }
}

document = Document()

document.add_heading('PDF Parser Results', 0)

document.add_heading('What surgeries has this patient had?', level=1)
for key in surgeryDict:
    document.add_paragraph(
        key, style='List Bullet'
    )
    dates = ', '.join(str(x) for x in surgeryDict[key]['datesList'])
    document.add_paragraph(
        'Date:    ' + dates, style='List Bullet 2'
    )
    pages = ', '.join(str(x) for x in surgeryDict[key]['pagesList'])
    document.add_paragraph(
        'Source pages:    ' + pages, style='List Bullet 2'
    )

document.add_heading('What medications has this patient used?', level=1)
for key in medicationDict:
    document.add_paragraph(
        key, style='List Bullet'
    )
    document.add_paragraph(
        'Start and end dates:    ', style='List Bullet 2'
    )
    for tup in medicationDict[key]['datesList']:
        document.add_paragraph(
            str(tup[0]) + ',  ' + str(tup[1]), style='List Bullet 3'
        )
    pages = ', '.join(str(x) for x in medicationDict[key]['pagesList'])
    document.add_paragraph(
        'Source pages:    ' + pages, style='List Bullet 2'
    )

document.add_heading('What allergies does the patient have?', level=1)
for key in allergyDict:
    document.add_paragraph(
        key, style='List Bullet'
    )
    pages = ', '.join(str(x) for x in allergyDict[key]['pagesList'])
    document.add_paragraph(
        'Source pages:    ' + pages, style='List Bullet 2'
    )

document.save('demo.docx')