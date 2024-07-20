import sys
from tempfile import TemporaryDirectory
from pathlib import Path
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import json
from docx import Document
from openai import OpenAI

client = OpenAI(api_key='')

PROMPT_SURGERIES = 'Using only this text, which surgeries are listed here, with their date, using None for unknown. If no surgeries return empty JSON. Use key of surgeries to a list of surgery key and date key: '
PROMPT_MEDICATIONS = 'Using only this text, which medications are listed here, with just their start date and end date, using None for unknown. If no medications return empty JSON. Use key of medications to a list of medication key and startDate key and endDate key: '
PROMPT_ALLERGIES = 'Using only this text, which allergies are listed here. If no allergies return empty JSON. Use key of allergies to a list of only allergyName key: '

surgeriesDict = {}
medicationsDict = {}
allergiesDict = {}

# RETURNS the dictionary of the API response to the prompt
def getDictionaryResponseFromAPI(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo-0125",
        response_format={ "type": "json_object" },
        messages=[
            {"role": "system", "content": "You are a helpful assistant designed to output JSON."},
            {"role": "user", "content": prompt}
        ]
    )
    respDict = json.loads(response.choices[0].message.content)
    return respDict


def addToSurgeriesDict(surgeryNameLower, surgeryDate, pageNum):
    sDate = surgeryDate
    if (sDate is None):
        sDate = 'None'

    if surgeryNameLower in surgeriesDict:
        if sDate != 'None':
            surgeriesDict[surgeryNameLower]['datesList'].append(sDate)
        surgeriesDict[surgeryNameLower]['pagesList'].append(pageNum)
    else:
        surgeriesDict[surgeryNameLower] = {
            'datesList': [sDate],
            'pagesList': [pageNum]
        }


def processSurgeriesDict(sDict, pageNum, text):
    if 'surgeries' in sDict:
        if isinstance(sDict['surgeries'], list):
            for s in sDict['surgeries']:
                if isinstance(s, dict):
                    if 'surgery' in s and 'date' in s:
                        if s['surgery'].lower() in text.lower():
                            addToSurgeriesDict(s['surgery'].lower(), s['date'], pageNum)


def addToMedicationsDict(medicationNameLower, startDate, endDate, pageNum):
    sDate = startDate
    if (sDate is None):
        sDate = 'None'
    
    eDate = endDate
    if (eDate is None):
        eDate = 'None'
    
    if medicationNameLower in medicationsDict:
        if sDate != 'None' or eDate != 'None':
            medicationsDict[medicationNameLower]['datesList'].append((sDate, eDate))
        medicationsDict[medicationNameLower]['pagesList'].append(pageNum)
    else:
        medicationsDict[medicationNameLower] = {
            'datesList': [(sDate, eDate)],
            'pagesList': [pageNum]
        }


def processMedicationsDict(mDict, pageNum, text):
    if 'medications' in mDict:
        if isinstance(mDict['medications'], list):
            for m in mDict['medications']:
                if isinstance(m, dict):
                    if 'medication' in m and 'startDate' in m and 'endDate' in m:
                        if m['medication'].lower() in text.lower():
                            addToMedicationsDict(m['medication'].lower(), m['startDate'], m['endDate'], pageNum)


def addToAllergiesDict(allergyNameLower, pageNum):
    if allergyNameLower in allergiesDict:
        allergiesDict[allergyNameLower]['pagesList'].append(pageNum)
    else:
        allergiesDict[allergyNameLower] = {
            'pagesList': [pageNum]
        }


def processAllergiesDict(aDict, pageNum, text):
    if 'allergies' in aDict:
        if isinstance(aDict['allergies'], list):
            for a in aDict['allergies']:
                if isinstance(a, dict):
                    if 'allergyName' in a:
                        if a['allergyName'].lower() in text.lower():
                            addToAllergiesDict(a['allergyName'].lower(), pageNum)


def printDicts():
    print()
    print(json.dumps(surgeriesDict, indent=4))
    print('------------------------------------------------------------')
    print(json.dumps(medicationsDict, indent=4))
    print('------------------------------------------------------------')
    print(json.dumps(allergiesDict, indent=4))
    print('------------------------------------------------------------')


def writeDicts():
    with open('program_dicts.txt', "a") as output_file:
        output_file.write(json.dumps(surgeriesDict, indent=4))
        output_file.write('------------------------------------------------------------')
        output_file.write(json.dumps(medicationsDict, indent=4))
        output_file.write('------------------------------------------------------------')
        output_file.write(json.dumps(allergiesDict, indent=4))


def writeToWordDoc():
    document = Document()

    document.add_heading('PDF Parser Results', 0)

    document.add_heading('What surgeries has this patient had?', level=1)
    for key in surgeriesDict:
        document.add_paragraph(
            key, style='List Bullet'
        )
        dates = ', '.join(str(x) for x in surgeriesDict[key]['datesList'])
        document.add_paragraph(
            'Date:    ' + dates, style='List Bullet 2'
        )
        pages = ', '.join(str(x) for x in surgeriesDict[key]['pagesList'])
        document.add_paragraph(
            'Source pages:    ' + pages, style='List Bullet 2'
        )

    document.add_heading('What medications has this patient used?', level=1)
    for key in medicationsDict:
        document.add_paragraph(
            key, style='List Bullet'
        )
        document.add_paragraph(
            'Start and end dates:    ', style='List Bullet 2'
        )
        for tup in medicationsDict[key]['datesList']:
            document.add_paragraph(
                str(tup[0]) + ',  ' + str(tup[1]), style='List Bullet 3'
            )
        pages = ', '.join(str(x) for x in medicationsDict[key]['pagesList'])
        document.add_paragraph(
            'Source pages:    ' + pages, style='List Bullet 2'
        )

    document.add_heading('What allergies does the patient have?', level=1)
    for key in allergiesDict:
        document.add_paragraph(
            key, style='List Bullet'
        )
        pages = ', '.join(str(x) for x in allergiesDict[key]['pagesList'])
        document.add_paragraph(
            'Source pages:    ' + pages, style='List Bullet 2'
        )

    document.save('PdfParserResults.docx')


def main(pdfName):
    outdir = Path(__file__).parent.resolve()
    pdfLocation = str(outdir) + '/' + pdfName
    
    image_file_list = []
    
    with TemporaryDirectory() as tempdir:
        pdf_pages = convert_from_path(pdfLocation, 500)
        
        for page_enumeration, page in enumerate(pdf_pages, start=1):
            filename = f"{tempdir}/page_{page_enumeration:03}.jpg"
            page.save(filename, "JPEG")
            image_file_list.append(filename)
        
        pageNumber = 1
        for image_file in image_file_list:
                text = str(((pytesseract.image_to_string(Image.open(image_file)))))
                text = text.replace("-\n", "")
                
                # Surgeries Section
                prompt = f"""{PROMPT_SURGERIES} {text}"""
                processSurgeriesDict(getDictionaryResponseFromAPI(prompt), pageNumber, text)
                
                # Medications Section
                prompt = f"""{PROMPT_MEDICATIONS} {text}"""
                processMedicationsDict(getDictionaryResponseFromAPI(prompt), pageNumber, text)
                
                # Allergies Section
                prompt = f"""{PROMPT_ALLERGIES} {text}"""
                processAllergiesDict(getDictionaryResponseFromAPI(prompt), pageNumber, text)
                
                pageNumber += 1
        
        writeToWordDoc()

if __name__ == "__main__":
    main(sys.argv[1])
